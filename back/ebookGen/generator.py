from docx.enum.section import WD_ORIENT
from docx.shared import Pt, RGBColor, Cm
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert
import PyPDF2

from coverPageGen import *
from pathlib import Path
import time

img_doc_path = "coverDoc/cover.pdf"
pdf_path = "pdf/myEbook.pdf"


def generate_file_path():
    current_time = time.strftime("%Y%m%d-%H%M%S")
    file_name = f"ebook_{current_time}.pdf"
    return f'generated_books/{file_name}'


def combine_pdfs(pdf_path1, pdf_path2, output_path):
    pdf_reader1 = PyPDF2.PdfReader(pdf_path1)
    pdf_reader2 = PyPDF2.PdfReader(pdf_path2)
    pdf_writer = PyPDF2.PdfWriter()

    for page_num in range(len(pdf_reader1.pages)):
        pdf_writer.add_page(pdf_reader1.pages[page_num])

    for page_num in range(len(pdf_reader2.pages)):
        pdf_writer.add_page(pdf_reader2.pages[page_num])

    with open(output_path, 'wb') as out_pdf_file:
        pdf_writer.write(out_pdf_file)


def create_word_book(book_info, content_array, filename):
    doc = Document()

    def setup_document_styles(doc):

        styles = [
            (
                'CoverTitleStyle', WD_STYLE_TYPE.PARAGRAPH, 'Georgia', 48, RGBColor(0, 0, 0),
                WD_PARAGRAPH_ALIGNMENT.CENTER,
                24, None),
            ('ChapterTitleStyle', WD_STYLE_TYPE.PARAGRAPH, 'Georgia', 26, RGBColor(0, 0, 0), None, 12, None),
            ('SubchapterTitleStyle', WD_STYLE_TYPE.PARAGRAPH, 'Georgia', 20, RGBColor(105, 105, 105), None, 8, None),
            ('ContentStyle', WD_STYLE_TYPE.PARAGRAPH, 'Times New Roman', 12, RGBColor(0, 0, 0), None, None, 1.5),
            ('TocChapterTitleStyle', WD_STYLE_TYPE.PARAGRAPH, 'Georgia', 18, RGBColor(0, 0, 0), None, 12, None),
            (
                'TOCHeadingStyle', WD_STYLE_TYPE.PARAGRAPH, 'Georgia', 24, RGBColor(0, 0, 0),
                WD_PARAGRAPH_ALIGNMENT.CENTER,
                12, None),
        ]

        for style_name, style_type, font_name, font_size, font_color, paragraph_alignment, space_after, line_spacing in styles:
            try:
                style = doc.styles[style_name]
            except KeyError:
                style = doc.styles.add_style(style_name, style_type)
            style.font.name = font_name
            style.font.size = Pt(font_size)
            style.font.color.rgb = font_color
            if paragraph_alignment is not None:
                style.paragraph_format.alignment = paragraph_alignment
            if space_after is not None:
                style.paragraph_format.space_after = Pt(space_after)
            if line_spacing is not None:
                style.paragraph_format.line_spacing = line_spacing

    def setup_document_layout(doc):
        """Sets up the page layout for the document."""
        section = doc.sections[0]
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_height = Cm(30)
        section.page_width = Cm(21)
        section.top_margin = Cm(2.75)
        section.bottom_margin = Cm(2.75)
        section.left_margin = Cm(2.75)
        section.right_margin = Cm(2.75)

    def add_cover_page(doc, book_info):
        doc.add_paragraph(book_info['ebook_title'], style='CoverTitleStyle')
        doc.add_page_break()

    def add_description_page(doc, book_info):
        doc.add_paragraph('description', style='TOCHeadingStyle')
        doc.add_paragraph(book_info['description'], style='ContentStyle')
        doc.add_page_break()

    def add_table_of_contents(doc, book_info):
        """Adds a table of contents pages."""
        doc.add_paragraph('Table of Contents', style='TOCHeadingStyle')

        for chapter_index, chapter in enumerate(book_info['chapters'], start=1):
            toc_text = f"{chapter['chapter_title']}"
            doc.add_paragraph(toc_text, style='TocChapterTitleStyle')

            for subsection_index, subsection_title in enumerate(chapter['subsections'], start=1):
                toc_subsection_text = f"    {chapter_index}.{subsection_index} {subsection_title}"
                para = doc.add_paragraph(toc_subsection_text, style='ContentStyle')
                para.paragraph_format.left_indent = Cm(2)

        doc.add_page_break()

    def add_content(doc, book_info, content_array):
        content_index = 0
        for chapter_index, chapter in enumerate(book_info['chapters'], start=1):
            chapT = f"{chapter['chapter_title']}"
            p = doc.add_paragraph(chapT, style='ChapterTitleStyle')
            p.add_run('\n' + '_' * 20).bold = True

            for subsection_index, subsection_title in enumerate(chapter['subsections'], start=1):
                subT = f"{chapter_index}.{subsection_index} {subsection_title}"

                doc.add_paragraph(subT, style='SubchapterTitleStyle')

                if content_index < len(content_array):
                    doc.add_paragraph(content_array[content_index], style='ContentStyle')
                    content_index += 1
                else:
                    doc.add_paragraph("Content missing", style='ContentStyle')
                doc.add_paragraph()
            doc.add_page_break()

    def add_summary(doc, book_info):
        doc.add_paragraph('summary', style='TOCHeadingStyle')
        doc.add_paragraph(book_info['summary'], style='ContentStyle')
        doc.add_page_break()

    def add_sources(doc, book_info):
        doc.add_paragraph('sources', style='TOCHeadingStyle')
        doc.add_paragraph(book_info['sources'], style='ContentStyle')
    # Setup document

    setup_document_styles(doc)
    setup_document_layout(doc)
    add_cover_page(doc, book_info)
    add_description_page(doc, book_info)
    add_table_of_contents(doc, book_info)
    add_content(doc, book_info, content_array)
    add_summary(doc, book_info)
    add_sources(doc, book_info)

    # Save the document
    doc_path = f"docs/{filename}.docx"
    doc.save(doc_path)
    convert(doc_path, pdf_path)


def create_ebook(filename, book_title, book_info, content_array, gender):
    path = generate_file_path()
    add_text_to_image(RESIZED_COVER_PATH, book_title, gender)
    create_word_book(book_info, content_array, filename)
    combine_pdfs(img_doc_path, pdf_path, path)

    return path
