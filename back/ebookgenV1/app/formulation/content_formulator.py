from docx.enum.section import WD_ORIENT
from docx.shared import Pt, RGBColor, Cm
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import Document
from docx2pdf import convert


class Style:

    def __init__(self, document: Document()):

        self.document = document

        self.styles = [
            ('CoverTitleStyle', WD_STYLE_TYPE.PARAGRAPH, 'Georgia', 48, RGBColor(0, 0, 0),
             WD_PARAGRAPH_ALIGNMENT.CENTER, 24, None),
            ('ChapterTitleStyle', WD_STYLE_TYPE.PARAGRAPH, 'Georgia', 26, RGBColor(0, 0, 0), None, 12, None),
            ('SubchapterTitleStyle', WD_STYLE_TYPE.PARAGRAPH, 'Georgia', 20, RGBColor(105, 105, 105), None, 8, None),
            ('ContentStyle', WD_STYLE_TYPE.PARAGRAPH, 'Times New Roman', 12, RGBColor(0, 0, 0), None, None, 1.5),
            ('TocChapterTitleStyle', WD_STYLE_TYPE.PARAGRAPH, 'Georgia', 18, RGBColor(0, 0, 0), None, 12, None),
            ('TOCHeadingStyle', WD_STYLE_TYPE.PARAGRAPH, 'Georgia', 24, RGBColor(0, 0, 0),
             WD_PARAGRAPH_ALIGNMENT.CENTER, 12, None),
        ]

    def setup_styles(self):
        """ Setting up the styles of the document"""
        for style_name, style_type, font_name, font_size, font_color, paragraph_alignment, space_after, line_spacing in self.styles:
            self._add_style(style_name, style_type, font_name, font_size, font_color,
                            paragraph_alignment, space_after, line_spacing)

    def _add_style(self, style_name, style_type, font_name, font_size, font_color,
                   paragraph_alignment, space_after, line_spacing):
        """
        Add or modify a style in the document.

        :param style_name: Name of the style.
        :param style_type: Type of the style (e.g., paragraph, character).
        :param font_name: Font name.
        :param font_size: Font size.
        :param font_color: Font color (RGBColor).
        :param paragraph_alignment: Paragraph alignment (e.g., CENTER, LEFT).
        :param space_after: Space after the paragraph (in points).
        :param line_spacing: Line spacing.
        """

        try:
            style = self.document.styles[style_name]
        except KeyError:
            style = self.document.styles.add_style(style_name, style_type)

        font = style.font
        font.name = font_name
        font.size = Pt(font_size)
        font.color.rgb = font_color

        paragraph_format = style.paragraph_format
        if paragraph_alignment is not None:
            paragraph_format.alignment = paragraph_alignment
        if space_after is not None:
            paragraph_format.space_after = Pt(space_after)
        if line_spacing is not None:
            paragraph_format.line_spacing = line_spacing


class ContentFormulation:

    def __init__(self,document: Document(), book_info=None):
        if book_info is None:
            book_info = {}
        self.book_info = book_info
        self.file_name = ""
        self.document = document

        book_style = Style(self.document)
        book_style.setup_styles()

    def setup_document_layout(self):
        """Sets up the page layout for the document."""
        section = self.document.sections[0]
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_height = Cm(30)
        section.page_width = Cm(21)
        section.top_margin = Cm(2.75)
        section.bottom_margin = Cm(2.75)
        section.left_margin = Cm(2.75)
        section.right_margin = Cm(2.75)

    def add_cover_page(self):
        self.document.paragraph(self.book_info['ebook_title'], style='CoverTitleStyle')
        self.document.page_break()

    def add_description_page(self):
        self.document.paragraph('description', style='TOCHeadingStyle')
        self.document.paragraph(self.book_info['description'], style='ContentStyle')
        self.document.page_break()

    def add_table_of_contents(self):
        """Adds a table of contents pages."""
        self.document.paragraph('Table of Contents', style='TOCHeadingStyle')

        for chapter_index, chapter in enumerate(self.book_info['chapters'], start=1):
            toc_text = f"{chapter['chapter_title']}"
            self.document.paragraph(toc_text, style='TocChapterTitleStyle')

            for subsection_index, subsection_title in enumerate(chapter['subsections'], start=1):
                toc_subsection_text = f"    {chapter_index}.{subsection_index} {subsection_title}"
                para = self.document.paragraph(toc_subsection_text, style='ContentStyle')
                para.paragraph_format.left_indent = Cm(2)

        self.document.page_break()

    def add_content(self, content_array):
        content_index = 0
        for chapter_index, chapter in enumerate(self.book_info['chapters'], start=1):
            chapT = f"{chapter['chapter_title']}"
            p = self.document.paragraph(chapT, style='ChapterTitleStyle')
            p.add_run('\n' + '_' * 20).bold = True

            for subsection_index, subsection_title in enumerate(chapter['subsections'], start=1):
                subT = f"{chapter_index}.{subsection_index} {subsection_title}"

                self.document.paragraph(subT, style='SubchapterTitleStyle')

                if content_index < len(content_array):
                    self.document.paragraph(content_array[content_index], style='ContentStyle')
                    content_index += 1
                else:
                    self.document.paragraph("Content missing", style='ContentStyle')
                self.document.paragraph()
            self.document.page_break()

    def add_summary(self):
        self.document.paragraph('summary', style='TOCHeadingStyle')
        self.document.paragraph(self.book_info['summary'], style='ContentStyle')
        self.document.page_break()

    def add_sources(self):
        self.document.paragraph('sources', style='TOCHeadingStyle')
        self.document.paragraph(self.book_info['sources'], style='ContentStyle')

    def save_document(self, file_name: str):
        self.file_name= f"docs/{file_name}.docx"



